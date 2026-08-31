import io
import re
import docx
import pandas as pd
import fitz  # PyMuPDF
from fastapi import UploadFile
import pytesseract
import os
import shutil  
import asyncio  
import gc
import threading
from PIL import Image
from concurrent.futures import ThreadPoolExecutor, as_completed

# --- STRICT CPU CONTROL FOR OCR LIBRARIES ---
os.environ["OMP_THREAD_LIMIT"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"
os.environ["OPENBLAS_NUM_THREADS"] = "1"

# --- ADVANCED EMBEDDED CONVERTERS ---
try:
    from docling.document_converter import DocumentConverter
    docling_converter = DocumentConverter()
    DOCLING_AVAILABLE = True
except ImportError:
    print("⚠️ WARNING: Docling library not detected. Falling back to native PyMuPDF.")
    DOCLING_AVAILABLE = False

# --- SMART TESSERACT CONFIGURATION ---
def configure_tesseract():
    linux_tesseract_path = shutil.which("tesseract")
    if linux_tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = linux_tesseract_path
        return

    appdata_path = os.path.join(os.path.expanduser('~'), 'AppData', 'Local', 'Programs', 'Tesseract-OCR', 'tesseract.exe')
    system_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

    if os.path.exists(appdata_path):
        pytesseract.pytesseract.tesseract_cmd = appdata_path
    elif os.path.exists(system_path):
        pytesseract.pytesseract.tesseract_cmd = system_path
    else:
        print("❌ ERROR: Tesseract OCR not found globally or in Windows paths.")

configure_tesseract()

def estimate_token_count(text: str) -> int:
    try:
        import tiktoken
        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))
    except Exception:
        return int(len(text.split()) * 1.37)

def clean_extracted_text(text: str) -> str:
    if not text: 
        return ""
        
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    text = re.sub(r'[ \t]{3,}', ' | ', text)
    text = re.sub(r'[ \t]{2}', ' ', text)
    text = re.sub(r'(?i)page\s*(?:no\.?)?\s*\d+\s*(?:of|/)?\s*\d*', '', text)
    text = re.sub(r'(?i)all\s+rights\s+reserved', '', text)
    text = re.sub(r'(?i)tender\s+document\s*(?:for)?', '', text)
    text = re.sub(r'(?i)commercial\s+bid\s+format', '', text)
    text = re.sub(r'(?i)strict\s+confidence\s*(?:confidential)?', '', text)
    text = re.sub(r'\d{2}[-/.]\d{2}[-/.]\d{4}\s+\d{2}:\d{2}(?::\d{2})?', '', text)
    text = re.sub(r'_{2,}', '', text)
    text = re.sub(r'-{3,}', '', text)
    text = re.sub(r'\.{3,}', '...', text) 
    
    lines = text.split('\n')
    seen_lines = []
    for line in lines:
        cleaned_line = line.strip()
        if cleaned_line and seen_lines and cleaned_line == seen_lines[-1]:
            continue
        seen_lines.append(cleaned_line)
    text = '\n'.join(seen_lines)
    
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'\n\|\s*\n', '\n', text)
    
    return text.strip()

async def extract_text_from_upload(file: UploadFile, task_id: str = None) -> str:
    print(f"\n--- [HYBRID START] Processing File: {file.filename} ---", flush=True)
    file_bytes = await file.read()
    
    raw_text = await asyncio.to_thread(extract_text_from_file, file_bytes, file.filename, task_id)
    cleaned_text = clean_extracted_text(raw_text)
    
    try:
        raw_token_estimate = estimate_token_count(raw_text)
        clean_token_estimate = estimate_token_count(cleaned_text)
        token_savings = raw_token_estimate - clean_token_estimate
        savings_percent = (token_savings / raw_token_estimate * 100) if raw_token_estimate > 0 else 0
        
        print("\n=======================================================", flush=True)
        print(f"📊 TASK METRICS ANALYSIS [{task_id or 'STANDALONE'}]", flush=True)
        print(f"  • Raw Character Length   : {len(raw_text)}", flush=True)
        print(f"  • Clean Character Length : {len(cleaned_text)}", flush=True)
        print(f"  • Estimated RAW Tokens   : {raw_token_estimate} tokens", flush=True)
        print(f"  • Estimated CLEAN Tokens : {clean_token_estimate} tokens", flush=True)
        print(f"  • 🔥 TOTAL TOKENS SAVED  : {token_savings} tokens ({savings_percent:.1f}% reduction)", flush=True)
        print("=======================================================\n", flush=True)
    except Exception as metrics_fault:
        print(f"[METRICS BYPASSED] Non-critical analytics fault encountered: {metrics_fault}", flush=True)
    
    print(f"--- [COMPLETE] Total Extracted Characters: {len(cleaned_text)} ---", flush=True)
    return cleaned_text

def extract_text_from_file(file_bytes: bytes, filename: str, task_id: str = None) -> str:
    """
    Memory-based extractor (Fallback/Legacy).
    Includes High-Speed OCR bypass mapping.
    """
    fn_lower = filename.lower()
    temp_pdf_path = f"temp_process_{task_id or 'standalone'}_{os.getpid()}.pdf"

    try:
        if fn_lower.endswith(".pdf"):
            with open(temp_pdf_path, "wb") as f:
                f.write(file_bytes)

            with fitz.open(temp_pdf_path) as doc:
                total_pages = len(doc)

            print(f"STATUS: PDF detected ({filename}). Processing {total_pages} Pages (High-Speed Memory Mode)...", flush=True)

            processed_count = 0
            progress_lock = threading.Lock()
            all_page_texts = [""] * total_pages

            BATCH_SIZE = 100
            max_workers = 3

            def process_single_page(page_num):
                try:
                    with fitz.open(temp_pdf_path) as local_doc:
                        page = local_doc[page_num]
                        current_page_display = page_num + 1
                        
                        # SPEED FIX: Extract text directly without expensive block arrays
                        extracted = page.get_text("text")
                        
                        if extracted and len(extracted.strip()) > 30:
                            page_text = f"\n--- Page {current_page_display} ---\n{extracted}\n"
                            log_msg = f"  > Page {current_page_display}/{total_pages}: PyMuPDF High-Speed Native Extraction"
                        else:
                            # SPEED FIX: Low DPI (1.0x matrix) renders ~40% faster
                            pix = page.get_pixmap(matrix=fitz.Matrix(1.0, 1.0))
                            img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("L")
                            ocr_result = pytesseract.image_to_string(img, lang="eng", config="--oem 1 --psm 6")
                            
                            img.close()
                            pix = None
                            
                            page_text = f"\n--- Page {current_page_display} (OCR Scan) ---\n{ocr_result}\n"
                            log_msg = f"  > Page {current_page_display}/{total_pages}: Tesseract OCR Fast Scan"

                        return page_num, page_text, log_msg
                except Exception as page_err:
                    return page_num, f"\n--- Page {page_num + 1} (Error) ---\n[Error: {str(page_err)}]\n", f"  ! Page {page_num + 1} Exception: {page_err}"

            for batch_start in range(0, total_pages, BATCH_SIZE):
                batch_end = min(batch_start + BATCH_SIZE, total_pages)
                
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = {executor.submit(process_single_page, i): i for i in range(batch_start, batch_end)}
                    
                    for future in as_completed(futures):
                        p_num, p_text, log_msg = future.result()
                        all_page_texts[p_num] = p_text
                        
                        print(log_msg, flush=True)
                        
                        with progress_lock:
                            processed_count += 1
                            if task_id:
                                try:
                                    from main import progress_store
                                    progress_store[task_id] = {"current": processed_count, "total": total_pages}
                                except ImportError:
                                    pass
                gc.collect()

            final_text = "".join(all_page_texts)
            all_page_texts.clear()
            gc.collect()

            return final_text

        elif fn_lower.endswith((".docx", ".doc")):
            doc_obj = docx.Document(io.BytesIO(file_bytes))
            text = ""
            for table in doc_obj.tables:
                for row in table.rows:
                    text += " | ".join([cell.text.strip() for cell in row.cells]) + "\n"
            text += "\n".join([p.text for p in doc_obj.paragraphs if p.text.strip()])
            return text

        elif fn_lower.endswith((".xlsx", ".xls", ".xlsm")):
            text = ""
            with pd.ExcelFile(io.BytesIO(file_bytes)) as xls:
                for sheet in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet)
                    if not df.empty:
                        text += f"\n[SHEET: {sheet}]\n{df.to_string(index=False)}\n"
            return text
            
        else: 
            return "Error: Unsupported file format."

    except Exception as e:
        print(f"!!! CRITICAL ERROR in extraction: {str(e)}", flush=True)
        return f"Error reading file {filename}: {str(e)}"
    finally:
        if os.path.exists(temp_pdf_path):
            try:
                os.remove(temp_pdf_path)
            except Exception:
                pass
        gc.collect()

def append_text_to_disk_stream(file_bytes: bytes, filename: str, output_file_path: str, task_id: str = None):
    """
    🚀 THE ULTIMATE UNLIMITED EXTRACTOR 🚀
    Streams extracted text DIRECTLY to the hard drive chunk-by-chunk.
    Uses High-Speed bypasses. Zero RAM crashes. Unlimited pages.
    """
    fn_lower = filename.lower()
    temp_pdf_path = f"temp_process_{task_id or 'standalone'}_{os.getpid()}.pdf"

    try:
        if fn_lower.endswith(".pdf"):
            with open(temp_pdf_path, "wb") as f:
                f.write(file_bytes)

            with fitz.open(temp_pdf_path) as doc:
                total_pages = len(doc)

            print(f"⚡ HIGH-SPEED DISK STREAM: Processing {total_pages} Pages...", flush=True)

            processed_count = 0
            progress_lock = threading.Lock()
            write_lock = threading.Lock()

            BATCH_SIZE = 100
            max_workers = 3 

            def process_single_page(page_num):
                try:
                    with fitz.open(temp_pdf_path) as local_doc:
                        page = local_doc[page_num]
                        current_page_display = page_num + 1
                        
                        # SPEED FIX: Direct fast text grab
                        extracted = page.get_text("text")
                        
                        if extracted and len(extracted.strip()) > 30:
                            page_text = f"\n--- Page {current_page_display} ---\n{extracted}\n"
                            log_msg = f"  > Page {current_page_display}/{total_pages}: PyMuPDF High-Speed Disk Write"
                        else:
                            # SPEED FIX: Low DPI (1.0x matrix)
                            pix = page.get_pixmap(matrix=fitz.Matrix(1.0, 1.0))
                            img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("L")
                            ocr_result = pytesseract.image_to_string(img, lang="eng", config="--oem 1 --psm 6")
                            img.close()
                            pix = None
                            page_text = f"\n--- Page {current_page_display} (OCR Scan) ---\n{ocr_result}\n"
                            log_msg = f"  > Page {current_page_display}/{total_pages}: Tesseract OCR Disk Write"

                        cleaned = clean_extracted_text(page_text)
                        
                        # Immediately flush to disk
                        with write_lock:
                            with open(output_file_path, "a", encoding="utf-8") as out_f:
                                out_f.write(cleaned + "\n")

                        return page_num, log_msg
                except Exception as err:
                    return page_num, f"  ! Page {page_num + 1} Error: {err}"

            for batch_start in range(0, total_pages, BATCH_SIZE):
                batch_end = min(batch_start + BATCH_SIZE, total_pages)
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = [executor.submit(process_single_page, i) for i in range(batch_start, batch_end)]
                    for future in as_completed(futures):
                        p_num, log_msg = future.result()
                        print(log_msg, flush=True)
                        
                        with progress_lock:
                            processed_count += 1
                            if task_id:
                                try:
                                    from main import progress_store
                                    progress_store[task_id] = {"current": processed_count, "total": total_pages}
                                except ImportError:
                                    pass
                gc.collect()

        elif fn_lower.endswith((".docx", ".doc")):
            doc_obj = docx.Document(io.BytesIO(file_bytes))
            with open(output_file_path, "a", encoding="utf-8") as out_f:
                for table in doc_obj.tables:
                    for row in table.rows:
                        out_f.write(" | ".join([cell.text.strip() for cell in row.cells]) + "\n")
                out_f.write("\n".join([p.text for p in doc_obj.paragraphs if p.text.strip()]) + "\n")

        elif fn_lower.endswith((".xlsx", ".xls", ".xlsm")):
            with pd.ExcelFile(io.BytesIO(file_bytes)) as xls:
                with open(output_file_path, "a", encoding="utf-8") as out_f:
                    for sheet in xls.sheet_names:
                        df = pd.read_excel(xls, sheet_name=sheet)
                        if not df.empty:
                            out_f.write(f"\n[SHEET: {sheet}]\n{df.to_string(index=False)}\n")

    except Exception as e:
        print(f"!!! ERROR in extraction stream: {str(e)}", flush=True)
    finally:
        if os.path.exists(temp_pdf_path):
            try:
                os.remove(temp_pdf_path)
            except Exception:
                pass
        gc.collect()